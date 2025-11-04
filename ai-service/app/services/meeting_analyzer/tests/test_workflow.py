"""
Teste manual do workflow (pode deletar depois).
"""
import base64
from app.workflow import meeting_analysis_app
from app.models.schema import MeetingAnalysisState


def create_test_docx_base64():
    """
    Cria um arquivo .docx de teste simulando uma TRANSCRIÇÃO REAL de reunião.
    Formato: [Nome]: fala
    """
    from docx import Document
    import io
    
    # Criar documento com transcrição realista
    doc = Document()
    doc.add_heading('Transcrição - Reunião Assessoria Financeira', 0)
    doc.add_paragraph('Data: 15/01/2025')
    doc.add_paragraph('Cliente: João Silva')
    doc.add_paragraph('Assessor: Maria Costa')
    doc.add_paragraph('Duração: 23 minutos')
    doc.add_paragraph('')
    doc.add_paragraph('---')
    doc.add_paragraph('')
    
    # Transcrição da conversa
    doc.add_paragraph('[Assessor Maria]: Bom dia João, como você está?')
    doc.add_paragraph('')
    doc.add_paragraph('[João Silva]: Bom dia Maria, tudo bem! Obrigado por me receber.')
    doc.add_paragraph('')
    doc.add_paragraph('[Assessor Maria]: Imagina! Então, o que te traz aqui hoje? Como posso te ajudar?')
    doc.add_paragraph('')
    doc.add_paragraph('[João Silva]: Olha, eu tenho 45 anos e estou começando a pensar mais seriamente sobre o futuro, sabe? Especialmente sobre aposentadoria. Quero me aposentar quando fizer 60 anos, mas não sei se estou no caminho certo.')
    doc.add_paragraph('')
    doc.add_paragraph('[Assessor Maria]: Entendo perfeitamente. E hoje, como está o seu patrimônio? Você já investe em alguma coisa?')
    doc.add_paragraph('')
    doc.add_paragraph('[João Silva]: Então, eu tenho uns 500 mil reais aplicados. Mas confesso que está tudo muito conservador... A maior parte tá na poupança mesmo, e um pouco em CDB. Eu sei que não é o ideal, mas tenho medo de perder dinheiro.')
    doc.add_paragraph('')
    doc.add_paragraph('[Assessor Maria]: Sem problemas, cada um tem seu perfil. Você se considera uma pessoa avessa a risco então?')
    doc.add_paragraph('')
    doc.add_paragraph('[João Silva]: Sim, bastante. Eu trabalho muito pra ganhar esse dinheiro, não quero ver ele desaparecer do dia pra noite. Prefiro ganhar menos mas dormir tranquilo.')
    doc.add_paragraph('')
    doc.add_paragraph('[Assessor Maria]: Perfeito, vou respeitar totalmente isso. E sobre a aposentadoria, você tem ideia de quanto gostaria de receber por mês quando se aposentar?')
    doc.add_paragraph('')
    doc.add_paragraph('[João Silva]: Olha, hoje eu ganho uns 35 mil por mês. Gostaria de manter um padrão parecido, mas se conseguir uns 25 a 30 mil já estaria ótimo.')
    doc.add_paragraph('')
    doc.add_paragraph('[Assessor Maria]: Entendi. E você já pensou em previdência privada? Tipo um PGBL ou VGBL?')
    doc.add_paragraph('')
    doc.add_paragraph('[João Silva]: Já ouvi falar, mas confesso que não entendo muito bem como funciona. Vale a pena? Você acha que eu deveria começar a fazer aportes? Quanto seria ideal?')
    doc.add_paragraph('')
    doc.add_paragraph('[Assessor Maria]: Acho que faz todo sentido pro seu caso. A gente pode fazer uma simulação, mas pensando em 15 anos até a aposentadoria, algo em torno de 3 mil por mês seria interessante.')
    doc.add_paragraph('')
    doc.add_paragraph('[João Silva]: 3 mil por mês eu consigo tranquilamente. Isso me daria uma boa aposentadoria?')
    doc.add_paragraph('')
    doc.add_paragraph('[Assessor Maria]: Sim, somado ao que você já tem investido, conseguimos estruturar algo bem sólido. Deixa eu te preparar uma proposta completa, pode ser?')
    doc.add_paragraph('')
    doc.add_paragraph('[João Silva]: Por favor, eu agradeço! Ah, outra coisa que me preocupa... Eu tenho dois filhos pequenos, 8 e 5 anos. Se acontecer algo comigo, eles ficam desprotegidos. Você trabalha com seguro de vida também?')
    doc.add_paragraph('')
    doc.add_paragraph('[Assessor Maria]: Trabalho sim! Isso é muito importante mesmo, ainda mais com filhos pequenos. Você já tem algum seguro hoje?')
    doc.add_paragraph('')
    doc.add_paragraph('[João Silva]: Não, nada. É algo que eu sei que deveria ter feito já, mas sempre fui adiando.')
    doc.add_paragraph('')
    doc.add_paragraph('[Assessor Maria]: Sem problemas, nunca é tarde. A gente pode estruturar uma cobertura adequada pro seu caso. Pensando na sua renda e nos seus filhos, algo em torno de 1 milhão de cobertura seria ideal.')
    doc.add_paragraph('')
    doc.add_paragraph('[João Silva]: Um milhão? Nossa, é caro?')
    doc.add_paragraph('')
    doc.add_paragraph('[Assessor Maria]: Na verdade não tanto quanto parece. Vou te fazer uma cotação e você vai se surpreender. Vale muito a pena pela tranquilidade.')
    doc.add_paragraph('')
    doc.add_paragraph('[João Silva]: Tá bom, manda pra mim. E sobre aquele dinheiro da poupança que eu te falei, tem como melhorar o rendimento sem correr muito risco?')
    doc.add_paragraph('')
    doc.add_paragraph('[Assessor Maria]: Com certeza! Existem fundos de renda fixa muito conservadores que rendem bem mais que a poupança. Vou te mostrar algumas opções que combinam com seu perfil.')
    doc.add_paragraph('')
    doc.add_paragraph('[João Silva]: Ótimo! Olha, acho que é isso que eu precisava mesmo. Me manda essas propostas que a gente conversa, ok?')
    doc.add_paragraph('')
    doc.add_paragraph('[Assessor Maria]: Perfeito João! Até quinta-feira eu te mando tudo: a simulação da previdência, a cotação do seguro e as opções de fundos. Aí a gente marca outra reunião pra detalhar tudo.')
    doc.add_paragraph('')
    doc.add_paragraph('[João Silva]: Combinado! Muito obrigado pela atenção, Maria.')
    doc.add_paragraph('')
    doc.add_paragraph('[Assessor Maria]: Por nada! Qualquer dúvida antes disso, pode me chamar. Boa semana!')
    doc.add_paragraph('')
    doc.add_paragraph('--- Fim da transcrição ---')
    
    # Salvar em memória
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    # Converter para base64
    file_bytes = file_stream.read()
    base64_data = base64.b64encode(file_bytes).decode('utf-8')
    
    # Formato data URI
    data_uri = f"data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{base64_data}"
    
    return data_uri


def test_workflow():
    """
    Testa o workflow completo.
    """
    print("=" * 60)
    print("TESTE DO WORKFLOW DE ANÁLISE DE REUNIÕES")
    print("=" * 60)
    
    # 1. Criar arquivo de teste
    print("\n1. Criando arquivo .docx de teste...")
    test_file = create_test_docx_base64()
    print(f"   ✅ Arquivo criado: {len(test_file)} caracteres (base64)")
    
    # 2. Criar estado inicial
    print("\n2. Criando estado inicial...")
    initial_state: MeetingAnalysisState = {
        "file_content": test_file,
        "file_name": "teste_reuniao.docx",
        "user_id": "test_user_123",
        "raw_text": "",
        "chunks": [],
        "partial_summaries": [],
        "partial_opportunities": [],
        "final_summary": "",
        "final_opportunities": [],
        "metadata": {},
        "error": None
    }
    print("   ✅ Estado inicial criado")
    
    # 3. Executar o workflow
    print("\n3. Executando workflow...\n")
    print("-" * 60)
    
    try:
        result = meeting_analysis_app.invoke(initial_state)
        
        print("-" * 60)
        print("\n4. Resultado:")
        print("=" * 60)
        
        # Verificar se houve erro
        if result.get('error'):
            print(f"❌ ERRO: {result['error']}")
            return
        
        # Mostrar resultados
        print(f"\n✅ Texto extraído:")
        print(f"   Tamanho: {len(result['raw_text'])} caracteres")
        print(f"   Preview: {result['raw_text'][:200]}...")
        
        print(f"\n✅ Chunks criados:")
        print(f"   Total: {len(result['chunks'])} chunks")
        for i, chunk in enumerate(result['chunks']):
            print(f"   Chunk {i+1}: {len(chunk)} caracteres")
            print(f"      Preview: {chunk[:100]}...")
        
        print(f"\n✅ Metadados:")
        for key, value in result['metadata'].items():
            print(f"   {key}: {value}")
        
        print("\n" + "=" * 60)
        print("TESTE CONCLUÍDO COM SUCESSO! 🎉")
        print("=" * 60)
        
    except Exception as e:
        print(f"\n❌ ERRO NA EXECUÇÃO: {str(e)}")
        import traceback
        traceback.print_exc()


if __name__ == "__main__":
    test_workflow()